"""Generate Word documents: Step-by-Step Guide and Business Primer."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x39, 0x64)   # headings
MID_BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # sub-headings
LIGHT_BLUE  = RGBColor(0xBD, 0xD7, 0xEE)   # table header fill
ACCENT      = RGBColor(0xED, 0x7D, 0x31)   # orange highlight
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
MED_GREY    = RGBColor(0xD9, 0xD9, 0xD9)


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = str(rgb)   # RGBColor.__str__ returns 'RRGGBB'
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        cfg = kwargs.get(side, {})
        el  = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   cfg.get("val",   "single"))
        el.set(qn("w:sz"),    cfg.get("sz",    "4"))
        el.set(qn("w:space"), cfg.get("space", "0"))
        el.set(qn("w:color"), cfg.get("color", "BFBFBF"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _row_bg(row, rgb: RGBColor):
    for cell in row.cells:
        _set_cell_bg(cell, rgb)


# ── Style helpers ─────────────────────────────────────────────────────────────

def _set_doc_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def _run(para, text, bold=False, italic=False, size=11,
         color=None, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(16)
    run.font.color.rgb = DARK_BLUE
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E75B6")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(13)
    run.font.color.rgb = MID_BLUE
    return p


def _heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK_BLUE
    return p


def _body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def _bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p


def _numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.add_run(text).font.size = Pt(11)
    return p


def _code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"),   "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"),  "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return p


def _note_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FFF2CC")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(f"  ⚠️  Note:  {text}")
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x40, 0x00)
    return p


def _table(doc, headers, rows, col_widths=None, zebra=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, DARK_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size  = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = LIGHT_GREY if (zebra and ri % 2 == 0) else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            if ci == 0:
                run = p.add_run(str(cell_text))
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_BLUE
            else:
                run = p.add_run(str(cell_text))
                run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            _set_col_width(t, i, w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def _title_page(doc, title, subtitle, version="v1.0  |  June 2026"):
    doc.add_paragraph().paragraph_format.space_before = Pt(60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size  = Pt(26)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size  = Pt(14)
    run2.font.color.rgb = MID_BLUE
    p2.paragraph_format.space_after = Pt(30)

    # Divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr  = div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E75B6")
    pBdr.append(bot)
    pPr.append(pBdr)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(version)
    r3.font.size  = Pt(10)
    r3.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    p3.paragraph_format.space_before = Pt(10)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 1 — STEP-BY-STEP TECHNICAL GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

def build_technical_guide():
    doc = Document()
    _set_doc_margins(doc)

    # ── Title page ────────────────────────────────────────────────────────────
    _title_page(
        doc,
        "RGM Trade Promotion ROI Optimizer",
        "Step-by-Step Application Guide",
    )

    # ══ Section 1: Overview ══════════════════════════════════════════════════
    _heading1(doc, "1.  Application Overview")
    _body(doc, (
        "The RGM Trade Promotion ROI Optimizer is a Streamlit-based AI co-pilot for "
        "Consumer Packaged Goods (CPG) Revenue Growth Management. It enables Trade "
        "Marketing Managers to plan, simulate, and compare quarterly promotional "
        "campaigns while enforcing compliance with all internal trade policies. "
        "The system uses Claude Sonnet (Anthropic) as an agentic planning engine "
        "with tool-use capabilities."
    ))

    _heading2(doc, "Technology Stack")
    _table(doc,
        ["Component", "Technology", "Version"],
        [
            ["UI Framework",     "Streamlit",         "≥ 1.35"],
            ["AI Engine",        "Claude Sonnet 4.6", "Anthropic SDK ≥ 0.40"],
            ["Data Processing",  "Pandas / NumPy",    "≥ 2.0 / 1.26"],
            ["Visualisation",    "Plotly",             "≥ 5.22"],
            ["Document Parsing", "python-docx",        "≥ 1.1"],
        ],
        col_widths=[4.5, 5.0, 3.0],
    )

    # ══ Section 2: Setup ══════════════════════════════════════════════════════
    _heading1(doc, "2.  Setup and Launch")
    _numbered(doc, "Clone or download the project to your local machine.")
    _numbered(doc, "Install all Python dependencies:")
    _code_block(doc, "pip install -r requirements.txt")
    _numbered(doc, "Launch the Streamlit application:")
    _code_block(doc, "streamlit run app.py")
    _numbered(doc, (
        "In the application sidebar, enter your ANTHROPIC_API_KEY. "
        "This key is required for the Claude agent to function."
    ))
    _numbered(doc, (
        "The app auto-loads all 7 CSV files from the data/ directory and parses "
        "business_rules.docx on startup. No manual file upload is needed."
    ))
    _note_box(doc, (
        "To use real trade data, replace the CSV files in data/ with your own files "
        "that match the schemas documented in Section 4."
    ))

    # ══ Section 3: Directory Structure ═══════════════════════════════════════
    _heading1(doc, "3.  Project Structure")
    _code_block(doc,
        "RGM_portfolio/\n"
        "├── app.py                    # Main Streamlit application\n"
        "├── requirements.txt\n"
        "├── business_rules.docx       # 17 compliance rules (parsed at runtime)\n"
        "├── agent/\n"
        "│   ├── agent_loop.py         # Claude agentic loop\n"
        "│   ├── tools.py              # 5 Claude tools + execution engine\n"
        "│   └── rules_parser.py       # Business rules DOCX parser\n"
        "├── analytics/\n"
        "│   ├── lift_model.py         # Price elasticity lift estimation\n"
        "│   ├── cannibalization.py    # Cross-SKU demand theft\n"
        "│   ├── roi_calculator.py     # ROI & financial metrics\n"
        "│   └── compliance_check.py   # 17-rule compliance validator\n"
        "└── data/\n"
        "    ├── generate_data.py      # Synthetic data generator\n"
        "    ├── calendar.csv\n"
        "    ├── sku_master.csv\n"
        "    ├── retailer_master.csv\n"
        "    ├── sales_history.csv\n"
        "    ├── promo_history.csv\n"
        "    ├── cannibalization_matrix.csv\n"
        "    └── trade_spend_ledger.csv"
    )

    # ══ Section 4: Input Files ════════════════════════════════════════════════
    _heading1(doc, "4.  Input Files and Field Schemas")
    _body(doc, (
        "The application reads seven CSV files. Each is described below with its "
        "full column schema and the downstream calculations that depend on each field."
    ))

    # --- 4.1 calendar.csv ---
    _heading2(doc, "4.1  calendar.csv  (104 rows — one row per week)")
    _body(doc, "Provides the time backbone for all analyses. Every other file links to this via week_id.")
    _table(doc,
        ["Column", "Type", "Example", "Role"],
        [
            ["week_id",          "string", "2025-W03",  "Primary time key — FK from all time-series files"],
            ["week_start_date",  "date",   "2025-01-13","Monday of the week"],
            ["week_end_date",    "date",   "2025-01-19","Sunday of the week"],
            ["month",            "int",    "1",          "Calendar month"],
            ["quarter",          "string", "Q1-2025",   "Fiscal quarter — primary planning unit"],
            ["fiscal_year",      "int",    "2025",       "Calendar year"],
            ["season",           "string", "Winter",    "Winter / Summer / Monsoon / Festive"],
            ["is_holiday_week",  "bool",   "True",       "Holiday flag"],
            ["festival_flag",    "string", "Diwali",    "Named festival or blank string"],
        ],
        col_widths=[3.5, 2.0, 3.0, 6.0],
    )
    _heading3(doc, "Key Field Interactions")
    _bullet(doc, "quarter  is the primary input for planning (e.g. 'Plan Q3-2025')")
    _bullet(doc, "season  feeds the SEASON_MULTIPLIER into lift calculations")
    _bullet(doc, "festival_flag ≠ ''  triggers FESTIVAL_MULTIPLIER = 1.40× and enforces Rule C02 (≥ 6 unique SKUs on promo)")

    # --- 4.2 sku_master.csv ---
    _heading2(doc, "4.2  sku_master.csv  (50 rows — one row per product)")
    _body(doc, "Master product catalogue. Drives eligibility checks, margin calculations, and elasticity lookups.")
    _table(doc,
        ["Column", "Type", "Example", "Role"],
        [
            ["sku_id",           "string", "SKU-0001",        "Primary key"],
            ["sku_name",         "string", "Premium Coffee 250g", "Display name"],
            ["brand",            "string", "BrewCo",           "Brand grouping (5 brands)"],
            ["category",         "string", "Coffee",           "Drives price elasticity lookup"],
            ["subcategory",      "string", "Premium Ground",   "Finer product grouping"],
            ["pack_size",        "string", "250g",             "Package quantity"],
            ["list_price",       "float",  "₹ 450.00",         "Undiscounted selling price"],
            ["cogs_per_unit",    "float",  "₹ 215.00",         "Cost of goods — used in GP calculation"],
            ["gross_margin_pct", "float",  "0.52",             "Gates BOGO eligibility (must be > 0.40)"],
            ["launch_date",      "date",   "2023-06-01",       "SKU must be ≥ 90 days old before promo"],
        ],
        col_widths=[3.5, 2.0, 3.5, 5.5],
    )
    _heading3(doc, "Key Field Interactions")
    _bullet(doc, "category → elasticity: Coffee=2.5, Tea=2.0, Juice=1.5, Chips=2.0, Biscuits=1.8")
    _bullet(doc, "list_price > ₹400  caps discount at 15% (Rule D01)")
    _bullet(doc, "gross_margin_pct ≤ 0.40  blocks BOGO mechanic (Rule M01)")
    _bullet(doc, "launch_date vs. promo start date  must be ≥ 90 days apart (Rule E01)")
    _bullet(doc, "brand + subcategory match  → high cannibalization (0.15–0.30 cross-elasticity)")

    # --- 4.3 retailer_master.csv ---
    _heading2(doc, "4.3  retailer_master.csv  (20 rows — one row per retailer)")
    _body(doc, "Trade partner directory. Tier classification drives most budget compliance rules.")
    _table(doc,
        ["Column", "Type", "Example", "Role"],
        [
            ["retailer_id",              "string", "RET-001",      "Primary key"],
            ["retailer_name",            "string", "BigMart",       "Display name"],
            ["channel",                  "string", "Modern Trade",  "Modern Trade / E-commerce / Traditional Trade"],
            ["region",                   "string", "North",         "Geographic segment"],
            ["tier",                     "string", "Tier 1",        "Tier 1 (5), Tier 2 (10), Tier 3 (5)"],
            ["coop_funding_pct",         "float",  "0.20",          "% of discount funded by retailer"],
            ["avg_weekly_volume_units",  "int",    "8500",          "Context for volume expectations"],
        ],
        col_widths=[3.8, 2.0, 3.0, 5.7],
    )
    _heading3(doc, "Key Field Interactions")
    _bullet(doc, "tier = 'Tier 1'  → ≥ 4 promos/quarter required (Rule C01); 35–55% of total budget (Rule B02)")
    _bullet(doc, "coop_funding_pct  → brand_investment = total_units × discount × list_price × (1 − coop_pct)")
    _bullet(doc, "Lift multiplier by tier: Tier 1 = 1.5×, Tier 2 = 0.8×, Tier 3 = 0.25×")

    # --- 4.4 sales_history.csv ---
    _heading2(doc, "4.4  sales_history.csv  (~97,000 rows)")
    _body(doc, "Weekly POS data. Non-promo weeks are averaged to produce future baseline forecasts.")
    _table(doc,
        ["Column", "Type", "Example", "Role"],
        [
            ["week_id",       "string", "2025-W03",   "FK to calendar"],
            ["sku_id",        "string", "SKU-0001",   "FK to sku_master"],
            ["retailer_id",   "string", "RET-001",    "FK to retailer_master"],
            ["units_sold",    "int",    "320",         "Actual units sold in the week"],
            ["gross_revenue", "float",  "₹ 144,000",  "Revenue for the week"],
            ["on_promo_flag", "bool",   "False",       "Whether this was a promo week"],
            ["baseline_units","int",    "210",         "Expected non-promo volume"],
            ["promo_id",      "string", "PROMO-0042", "FK to promo_history; blank if not on promo"],
        ],
        col_widths=[3.5, 2.0, 3.0, 6.0],
    )
    _heading3(doc, "Key Field Interactions")
    _bullet(doc, "Rows where on_promo_flag = False  are averaged to compute projected_weekly_baseline_units for planning")
    _bullet(doc, "Falls back to the same quarter in the prior fiscal year if current quarter data is sparse")
    _bullet(doc, "units_sold − baseline_units  (when on_promo_flag = True) shows the historically realized lift")

    # --- 4.5 promo_history.csv ---
    _heading2(doc, "4.5  promo_history.csv  (800 rows — past campaigns)")
    _body(doc, "Historical promotion events. Used exclusively for the 13-week rolling window compliance check.")
    _table(doc,
        ["Column", "Type", "Example", "Role"],
        [
            ["promo_id",          "string", "PROMO-0042", "Primary key"],
            ["sku_id",            "string", "SKU-0001",   "Product promoted"],
            ["retailer_id",       "string", "RET-001",    "Where it ran"],
            ["start_week",        "string", "2025-W01",   "Week promo launched"],
            ["end_week",          "string", "2025-W02",   "Week promo ended (inclusive)"],
            ["mechanic",          "string", "Price Off",  "Promotional tactic"],
            ["discount_depth_pct","float",  "0.15",       "% off from list_price"],
            ["actual_units",      "int",    "1,100",      "Realized units sold"],
            ["realized_lift_pct", "float",  "0.82",       "Actual lift vs. baseline"],
        ],
        col_widths=[3.5, 2.0, 3.0, 6.0],
    )
    _heading3(doc, "Key Field Interactions")
    _bullet(doc, "start_week / end_week  → counted in the 13-week rolling window for Rule C03 (≤ 10 promo weeks per window)")
    _bullet(doc, "mechanic  multipliers: Price Off=1.0×, BOGO=1.8×, Display=0.6×, Feature=0.8×, Combo=0.7×")

    # --- 4.6 cannibalization_matrix.csv ---
    _heading2(doc, "4.6  cannibalization_matrix.csv  (450 rows — SKU pairs)")
    _body(doc, "Cross-elasticity matrix quantifying demand theft between simultaneously promoted products at the same retailer.")
    _table(doc,
        ["Column", "Type", "Example", "Role"],
        [
            ["sku_a",            "string", "SKU-0001", "The SKU being promoted"],
            ["sku_b",            "string", "SKU-0002", "The SKU losing share"],
            ["cross_elasticity", "float",  "0.22",     "% of SKU-A's incremental units stolen from SKU-B"],
            ["relationship_type","string", "Same brand","Same brand vs. Competing brand"],
        ],
        col_widths=[3.5, 2.0, 2.5, 6.5],
    )
    _heading3(doc, "Key Field Interactions")
    _bullet(doc, "Applied when sku_a and sku_b are promoted at the same retailer in overlapping weeks")
    _bullet(doc, "net_incremental_units(sku_a) = gross_incremental − Σ(incremental × cross_elasticity for each simultaneous sku_b)")
    _bullet(doc, "cross_elasticity > 0.20  between simultaneous promos triggers Rule X01 violation")
    _bullet(doc, "The matrix is asymmetric: theft A→B ≠ theft B→A")

    # --- 4.7 trade_spend_ledger.csv ---
    _heading2(doc, "4.7  trade_spend_ledger.csv  (~2,000 rows — accounting transactions)")
    _body(doc, "GL-level trade spend transactions for post-hoc audit and reconciliation. Not used in live planning calculations.")
    _table(doc,
        ["Column", "Type", "Example", "Role"],
        [
            ["transaction_id", "string", "TXN-00123",   "Unique ledger entry"],
            ["promo_id",       "string", "PROMO-0042",  "FK to promo_history"],
            ["retailer_id",    "string", "RET-001",     "Retailer invoiced"],
            ["spend_type",     "string", "Off-Invoice", "Off-Invoice / Scan-Back / Slotting / Display Fee / Co-op"],
            ["amount",         "float",  "₹ 3,500",     "Transaction amount"],
            ["posted_date",    "date",   "2025-03-15",  "Accounting date"],
            ["period",         "string", "2025-Q1",     "Fiscal period for reporting"],
        ],
        col_widths=[3.5, 2.0, 3.0, 6.0],
    )

    # ══ Section 5: Processing Pipeline ═══════════════════════════════════════
    _heading1(doc, "5.  Processing Pipeline")

    _heading2(doc, "5.1  Baseline Forecast  (Tool: get_baseline_forecast)")
    _bullet(doc, "Input:  sku_id, retailer_id, quarter")
    _bullet(doc, "Queries sales_history for rows where on_promo_flag = False in the target quarter")
    _bullet(doc, "Falls back to the same quarter in the prior year if data is sparse")
    _bullet(doc, "Output:  projected_weekly_baseline_units, yoy_trend")

    _heading2(doc, "5.2  Lift Estimation  (Tool: estimate_promo_lift)")
    _body(doc, "Core formula:")
    _code_block(doc, "lift_pct = elasticity(category) × discount_pct × mechanic_multiplier × seasonal_multiplier")
    _body(doc, "Multiplier lookup tables:")
    _table(doc,
        ["Factor", "Source", "Values"],
        [
            ["Category elasticity",  "sku_master.category",      "Coffee=2.5, Tea=2.0, Juice=1.5, Chips=2.0, Biscuits=1.8"],
            ["Mechanic multiplier",  "promo event.mechanic",     "Price Off=1.0, BOGO=1.8, Display=0.6, Feature=0.8, Combo=0.7"],
            ["Season multiplier",    "calendar.season",          "Winter=0.9, Summer=1.1, Monsoon=0.8, Festive=1.3"],
            ["Festival multiplier",  "calendar.festival_flag",   "1.40× overrides season when festival_flag ≠ ''"],
            ["Tier multiplier",      "retailer_master.tier",     "Tier 1=1.5×, Tier 2=0.8×, Tier 3=0.25×"],
        ],
        col_widths=[4.0, 4.5, 6.0],
    )

    _heading2(doc, "5.3  Cannibalization Adjustment  (Tool: simulate_plan)")
    _code_block(doc,
        "for each event A:\n"
        "    for each simultaneous event B at same retailer:\n"
        "        loss += A.incremental_units × cross_elasticity(A, B)\n"
        "net_incremental_units(A) = max(0, A.incremental_units − loss)"
    )

    _heading2(doc, "5.4  ROI Calculation")
    _code_block(doc,
        "promo_price             = list_price × (1 − discount_depth_pct)\n"
        "incremental_revenue     = net_incremental_units × promo_price\n"
        "incremental_gp          = net_incremental_units × (promo_price − cogs_per_unit)\n"
        "brand_investment        = total_promo_units × list_price × discount × (1 − coop_pct)\n"
        "event_roi               = incremental_gp / brand_investment\n"
        "plan_roi                = Σ(incremental_gp) / Σ(brand_investment)"
    )

    # ══ Section 6: Compliance Rules ══════════════════════════════════════════
    _heading1(doc, "6.  Compliance Rules Reference")
    _body(doc, "All 17 rules are checked automatically during simulate_plan. Red = violation blocks submission.")
    _table(doc,
        ["Rule ID", "Category", "Condition", "Action if Violated"],
        [
            ["B01", "Budget",      "Total plan spend > ₹50M",                              "Reduce events or discount depth"],
            ["B02", "Budget",      "Tier 1 spend < 35% or > 55% of total",                 "Rebalance spend between tiers"],
            ["B03", "Budget",      "Single SKU > 12% of total spend",                      "Spread budget across more SKUs"],
            ["D01", "Discount",    "list_price > ₹400 AND discount > 15%",                 "Reduce discount to ≤ 15%"],
            ["D02", "Discount",    "Any discount > 30%",                                   "Reduce discount to ≤ 30%"],
            ["D03", "Discount",    "Any discount < 5%",                                    "Raise discount to ≥ 5%"],
            ["M01", "Mechanic",    "BOGO on SKU with gross_margin_pct ≤ 0.40",            "Switch mechanic or choose higher-margin SKU"],
            ["M02", "Mechanic",    "Combo with < 2 subcategories",                         "Add SKU from different subcategory"],
            ["M03", "Mechanic",    "Display event < 2 weeks duration",                     "Extend to ≥ 2 weeks"],
            ["E01", "Eligibility", "SKU launched < 90 days before promo start",            "Choose a more mature SKU"],
            ["C01", "Coverage",    "Tier 1 retailer with < 4 promos in quarter",           "Add events at that retailer"],
            ["C02", "Coverage",    "Festival week with < 6 unique SKUs on promo",          "Add SKUs during festival week"],
            ["C03", "Coverage",    "SKU at retailer on promo > 10 of any 13 weeks",        "Shorten or remove event"],
            ["X01", "Cannibalism", "Simultaneous promos with cross_elasticity > 0.20",     "Schedule in different weeks or swap SKU"],
            ["R01", "Compliance",  "Regional compliance flag",                             "Review regional guidelines"],
            ["R02", "Compliance",  "Channel compliance flag",                              "Review channel guidelines"],
        ],
        col_widths=[2.0, 2.5, 5.5, 4.5],
    )

    # ══ Section 7: UI Tabs ════════════════════════════════════════════════════
    _heading1(doc, "7.  Application Screens")

    _heading2(doc, "7.1  Plan Builder (Tab 1)")
    _bullet(doc, "Chat interface — type natural language requests to Claude")
    _bullet(doc, "Tool trace panel shows each tool call, input, output summary, and duration")
    _bullet(doc, "Saved scenarios are displayed in a quick-view sidebar panel")
    _bullet(doc, "Recommended prompt format:  'Plan [quarter] promotions for [brand] within ₹[X]M budget'")

    _heading2(doc, "7.2  Scenario Compare (Tab 2)")
    _bullet(doc, "Multi-select up to 4 saved scenarios")
    _bullet(doc, "Metrics table: ROI, total spend, incremental revenue, incremental GP, compliance status")
    _bullet(doc, "ROI bar chart and spend vs. GP waterfall chart")
    _bullet(doc, "Per-rule compliance grid and per-event breakdown table")

    _heading2(doc, "7.3  Rationale Inspector (Tab 3)")
    _bullet(doc, "Select scenario → select event → full audit trail rendered")
    _bullet(doc, "Sections: baseline forecast, lift multiplier breakdown, cannibalization deductions, ROI math, rules pass/fail")
    _bullet(doc, "ROI waterfall chart for the selected event")

    # ══ Section 8: Typical Workflow ══════════════════════════════════════════
    _heading1(doc, "8.  Typical End-to-End Workflow")
    _numbered(doc, "Launch app and enter API key in sidebar.")
    _numbered(doc, "Go to Plan Builder. Type your planning brief (brand, quarter, budget, any constraints).")
    _numbered(doc, "Review AI-generated plan summary: ROI, spend, and compliance status.")
    _numbered(doc, "If violations are flagged, instruct the AI to fix them in natural language.")
    _numbered(doc, "Save compliant plan as Scenario-001.")
    _numbered(doc, "Request an alternative plan (different mechanics, budget split, or SKU selection).")
    _numbered(doc, "Save as Scenario-002.")
    _numbered(doc, "Switch to Scenario Compare tab. Select both scenarios. Identify the better performer.")
    _numbered(doc, "Switch to Rationale Inspector. Audit the highest-spend event for any selected scenario.")
    _numbered(doc, "Export plan summary for internal review or submission.")

    # ══ Section 9: Quick Reference ════════════════════════════════════════════
    _heading1(doc, "9.  Quick Reference — Which File Drives What")
    _table(doc,
        ["Business Action", "File to Update", "Key Column"],
        [
            ["Add a new product",                       "sku_master.csv",            "sku_id, list_price, cogs_per_unit, launch_date"],
            ["Add a new retailer",                      "retailer_master.csv",       "retailer_id, tier, coop_funding_pct"],
            ["Change a product's price or margin",      "sku_master.csv",            "list_price, cogs_per_unit, gross_margin_pct"],
            ["Change retailer co-funding share",        "retailer_master.csv",       "coop_funding_pct"],
            ["Adjust seasonal lift multipliers",        "calendar.csv",              "season, festival_flag"],
            ["Update historical sales baselines",       "sales_history.csv",         "units_sold, baseline_units, on_promo_flag"],
            ["Record a past promotion (C03 window)",    "promo_history.csv",         "start_week, end_week, sku_id, retailer_id"],
            ["Update cross-SKU cannibalization rates",  "cannibalization_matrix.csv","cross_elasticity"],
            ["Record trade spend for reconciliation",   "trade_spend_ledger.csv",    "amount, promo_id, period"],
        ],
        col_widths=[5.0, 4.5, 5.0],
    )

    doc.save("RGM_Step_by_Step_Guide.docx")
    print("Saved: RGM_Step_by_Step_Guide.docx")


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 2 — BUSINESS USER PRIMER
# ═══════════════════════════════════════════════════════════════════════════════

def build_business_primer():
    doc = Document()
    _set_doc_margins(doc)

    # ── Title page ────────────────────────────────────────────────────────────
    _title_page(
        doc,
        "Trade Promotion ROI Optimizer",
        "Business User Primer",
    )

    # ══ What Is This Tool ════════════════════════════════════════════════════
    _heading1(doc, "What Is This Tool?")
    _body(doc, (
        "The Trade Promotion ROI Optimizer is your AI-powered co-pilot for planning "
        "quarterly promotional campaigns. Instead of building complex spreadsheets, you "
        "describe what you want to achieve in plain English and the AI designs, tests, "
        "and refines a promotion plan for you — in minutes."
    ))
    _body(doc, "The tool answers three key business questions before you spend a single rupee:")
    _bullet(doc, "Will this plan make money?",          bold_prefix="")
    _bullet(doc, "Does it meet our trade spend policies?", bold_prefix="")
    _bullet(doc, "Which of my options is the best?",    bold_prefix="")

    p = doc.paragraphs[-1]

    # ══ Getting Started ══════════════════════════════════════════════════════
    _heading1(doc, "Getting Started")
    _body(doc, (
        "The only thing you need is your API key, which your IT or analytics team will "
        "provide. Enter it in the sidebar when the app opens. Everything else — product "
        "data, retailer data, historical sales, past promotions — is already loaded. "
        "You just talk to the AI."
    ))

    # ══ Three Screens ════════════════════════════════════════════════════════
    _heading1(doc, "The Three Screens")
    _body(doc, "The app has three tabs. Think of them as a natural left-to-right workflow.")
    _table(doc,
        ["Screen", "Tab Name", "What You Do Here"],
        [
            ["1", "Plan Builder",         "Build your promotion plan by chatting with the AI"],
            ["2", "Scenario Compare",     "Compare two or more plan options side by side"],
            ["3", "Rationale Inspector",  "Audit any single promotional event in detail"],
        ],
        col_widths=[1.5, 4.0, 9.0],
    )

    # ══ Screen 1: Plan Builder ════════════════════════════════════════════════
    _heading1(doc, "Screen 1 — Plan Builder:  Where You Work with the AI")
    _body(doc, (
        "This is a chat window. You type your planning request, the AI builds a plan, "
        "and you refine it together through natural conversation."
    ))

    _heading2(doc, "How to Start a Conversation")
    _body(doc, "Be specific about three things:  brand or product,  quarter,  and  budget.")
    _code_block(doc, '"Plan Q3-2025 promotions for BrewCo coffee products. My total trade budget is ₹20 million."')
    _body(doc, "That is all you need. The AI will take it from there.")

    _heading2(doc, "What the AI Does Behind the Scenes")
    _body(doc, "You do not need to configure any of this — it happens automatically.")
    _table(doc,
        ["Step", "What the AI Is Doing", "Why It Matters to You"],
        [
            ["1", "Looks up your products, prices, and margins",
             "Knows which SKUs are eligible and how profitable each promotion will be"],
            ["2", "Pulls 2 years of weekly sales data",
             "Establishes a realistic sales baseline before the promotion starts"],
            ["3", "Estimates how many extra units each promo will sell",
             "Accounts for the discount size, promotion type, season, and upcoming festivals"],
            ["4", "Checks whether similar products promoted simultaneously will steal sales from each other",
             "Prevents you from running events that cancel each other out"],
            ["5", "Calculates your investment and return for every single event",
             "Tells you exactly what each rupee of trade spend earns back"],
            ["6", "Runs 17 compliance checks automatically",
             "Flags policy violations before the plan is submitted — no manual checking needed"],
        ],
        col_widths=[1.0, 5.5, 8.0],
    )

    _heading2(doc, "Reading the Plan Output")
    _table(doc,
        ["Metric", "What It Means"],
        [
            ["Total Brand Spend",    "How much you are investing in trade after the retailer's co-funding contribution"],
            ["Incremental Revenue",  "Extra revenue generated above your normal baseline by the promotions"],
            ["Incremental Gross Profit", "Extra profit after subtracting the cost of goods sold"],
            ["Plan ROI",             "For every ₹1 you invest, how much gross profit comes back. ROI of 1.5 = ₹1.50 returned per ₹1 spent"],
            ["Compliance Status",    "Green = plan is clean and ready to submit. Red = violations that must be fixed first"],
        ],
        col_widths=[4.5, 10.0],
    )

    _heading2(doc, "Refining Your Plan")
    _body(doc, (
        "If the AI flags a compliance violation, simply tell it what to fix in plain "
        "language. You do not need to understand the underlying numbers."
    ))
    _code_block(doc, '"The Tier 1 budget allocation is too low. Shift more spend to BigMart and SuperShop."')
    _code_block(doc, '"Remove the BOGO on the premium coffee SKU — swap it for a Price Off instead."')
    _code_block(doc, '"Bring all discounts on products above ₹400 down to 12%."')
    _body(doc, "The AI will adjust and re-simulate instantly.")

    _heading2(doc, "Saving a Scenario")
    _body(doc, (
        "Every simulated plan is saved automatically. Run as many variations as you like — "
        "conservative budget, aggressive budget, different mechanics — then compare them "
        "on the next screen."
    ))

    # ══ Compliance Rules ══════════════════════════════════════════════════════
    _heading1(doc, "The 17 Compliance Rules — Plain English")
    _body(doc, (
        "The tool automatically checks all 17 trade policy rules every time a plan is "
        "simulated. You never need to check these yourself. Below is what each rule "
        "means in business terms."
    ))

    _heading2(doc, "Budget Rules")
    _table(doc,
        ["Rule", "Plain English", "Common Fix"],
        [
            ["Total budget cap",         "Your total quarterly trade spend cannot exceed ₹50 million",
             "Reduce number of events or discount depths"],
            ["Tier 1 allocation",        "35%–55% of budget must go to your top-tier retailers. Not too little, not too much.",
             "Shift spend between Tier 1 and other retailers"],
            ["Single-SKU concentration", "No one product can consume more than 12% of your total budget",
             "Spread budget across more products"],
        ],
        col_widths=[4.0, 6.5, 4.0],
    )

    _heading2(doc, "Discount Rules")
    _table(doc,
        ["Rule", "Plain English", "Common Fix"],
        [
            ["Premium product cap", "Products priced above ₹400 can be discounted no more than 15%.\nDeep discounts on premium SKUs train consumers to wait for deals.",
             "Reduce discount to ≤ 15%"],
            ["Maximum discount",    "No promotion can offer more than 30% off list price",
             "Reduce discount to ≤ 30%"],
            ["Minimum discount",    "Promotions below 5% off are not meaningful and will not be approved",
             "Raise discount to ≥ 5%"],
        ],
        col_widths=[4.0, 6.5, 4.0],
    )

    _heading2(doc, "Promotion Mechanic Rules")
    _table(doc,
        ["Rule", "Plain English", "Common Fix"],
        [
            ["BOGO eligibility",  "Buy-One-Get-One deals are only allowed on products with a gross margin above 40%.\nOtherwise the deal loses money.",
             "Switch to Price Off or choose a higher-margin SKU"],
            ["Combo pack rule",   "A combo deal must bundle products from at least 2 different subcategories",
             "Add a SKU from a different subcategory to the bundle"],
            ["Display minimum",   "In-store display events must run for at least 2 weeks to justify the setup cost",
             "Extend the event duration"],
        ],
        col_widths=[4.0, 6.5, 4.0],
    )

    _heading2(doc, "Coverage Rules")
    _table(doc,
        ["Rule", "Plain English", "Common Fix"],
        [
            ["Tier 1 frequency",    "Each of your top-tier retailers must have at least 4 promotional events in the quarter",
             "Add more events at that retailer"],
            ["Festival coverage",   "During a major festival week (Diwali, Holi, etc.), at least 6 different products must be on promotion simultaneously",
             "Add products during the festival week"],
            ["Promo fatigue",       "The same product at the same retailer cannot be on promotion for more than 10 out of any 13 consecutive weeks.\nThis protects your sales baseline.",
             "Shorten the event or remove it"],
        ],
        col_widths=[4.0, 6.5, 4.0],
    )

    _heading2(doc, "Eligibility and Conflict Rules")
    _table(doc,
        ["Rule", "Plain English", "Common Fix"],
        [
            ["New product rule",        "A product must have been on shelf for at least 90 days before it can be promoted.\nThis ensures a real sales baseline exists.",
             "Choose a more established product"],
            ["Cannibalization conflict", "If two products at the same store are so closely related that promoting both simultaneously causes them to heavily steal sales from each other, they cannot run at the same time.",
             "Stagger the events across different weeks, or swap one SKU for an unrelated product"],
        ],
        col_widths=[4.0, 6.5, 4.0],
    )

    # ══ Screen 2 ══════════════════════════════════════════════════════════════
    _heading1(doc, "Screen 2 — Scenario Compare:  Pick Your Best Plan")
    _numbered(doc, "Select up to 4 saved scenarios from the dropdown.")
    _numbered(doc, "Review the side-by-side comparison table.")
    _numbered(doc, "Use the ROI bar chart and waterfall chart to visualise the difference.")
    _numbered(doc, "Drill into per-event detail to understand what is driving the difference.")

    _body(doc, "")
    _body(doc, "How to choose between scenarios:")
    _table(doc,
        ["Your Priority", "Pick the Scenario With..."],
        [
            ["Maximum total profit",         "Highest Incremental Gross Profit"],
            ["Best capital efficiency",       "Highest Plan ROI"],
            ["Staying within a budget cap",   "Total Brand Spend within your limit — and highest ROI within that"],
            ["Safest submission",             "Compliant = Yes, then highest ROI"],
        ],
        col_widths=[5.5, 9.0],
    )

    # ══ Screen 3 ══════════════════════════════════════════════════════════════
    _heading1(doc, "Screen 3 — Rationale Inspector:  Understand Any Single Event")
    _body(doc, (
        "Before presenting a plan to leadership, open this tab and pull up the full "
        "audit trail for any individual promotional event."
    ))
    _body(doc, "The inspector shows you:")
    _table(doc,
        ["Section", "What It Tells You"],
        [
            ["Baseline",         "What this product normally sells at this retailer in a non-promo week, based on 2 years of history"],
            ["Lift Estimate",    "How many extra units the promotion is expected to sell, and why — broken down by discount size, promotion type, and season"],
            ["Cannibalization",  "Whether any units are being deducted because a closely related product is running at the same store at the same time"],
            ["ROI Math",         "Exact calculation: extra revenue, extra gross profit, your cash investment after retailer co-funding"],
            ["Rules",            "Which of the 17 rules apply to this specific event and whether each one passes or fails"],
        ],
        col_widths=[3.5, 11.0],
    )
    _note_box(doc, "Use the Rationale Inspector before any senior review meeting. If someone asks 'Why did you choose a 20% discount at BigMart?', this tab gives you the precise answer.")

    # ══ Typical Session ══════════════════════════════════════════════════════
    _heading1(doc, "A Typical Session from Start to Finish")
    _table(doc,
        ["Time", "Action", "What Happens"],
        [
            ["10:00 AM", "Open app, enter API key, go to Plan Builder",
             "App loads with all data pre-populated"],
            ["10:02 AM", "Type planning brief — brand, quarter, budget, any constraints",
             "AI drafts an initial plan with ~10–15 promotional events"],
            ["10:03 AM", "Review ROI summary and compliance status",
             "If violations exist, AI explains which rule was triggered and why"],
            ["10:05 AM", "Instruct AI to fix violations in plain English",
             "AI adjusts and re-simulates the plan instantly"],
            ["10:06 AM", "Save compliant plan as Scenario-001",
             "Plan is stored for later comparison"],
            ["10:07 AM", "Request a second variation (e.g. different mechanics or budget split)",
             "AI generates Scenario-002"],
            ["10:10 AM", "Switch to Scenario Compare — select both scenarios",
             "Side-by-side metrics table and charts appear"],
            ["10:12 AM", "Pick the better-performing scenario",
             "Identify which plan delivers more profit or better ROI"],
            ["10:14 AM", "Open Rationale Inspector — audit the biggest-spend event",
             "Full breakdown confirms numbers are defensible"],
            ["10:15 AM", "Export plan for presentation",
             "Plan is ready for internal review or budget submission"],
        ],
        col_widths=[2.0, 5.0, 7.5],
    )

    # ══ Tips ══════════════════════════════════════════════════════════════════
    _heading1(doc, "Tips for Getting the Best Results")
    _bullet(doc, (
        "Be specific about constraints.  The more context you give, the better the first draft.\n"
        'Instead of "Plan some promos for Q3", say:\n'
        '"Plan Q3-2025 promos for BrewCo and SipWell only, no more than ₹20M total, '
        'avoid BOGO mechanics, prioritise Modern Trade channel."'
    ))
    _bullet(doc, (
        "Let the AI handle compliance.  Do not try to pre-calculate which rules apply. "
        "Just tell the AI your business intent and let it flag violations."
    ))
    _bullet(doc, (
        "Run at least two scenarios.  A conservative and an aggressive version gives you "
        "a defensible range for budget conversations."
    ))
    _bullet(doc, (
        "Use the Rationale Inspector before any senior review.  If someone asks 'why did "
        "you choose a 20% discount at BigMart?', the inspector gives you the exact answer "
        "in three bullet points."
    ))
    _bullet(doc, (
        "If the AI misunderstands, correct it directly.  "
        '"No — I meant a 15% discount, not 25%. Please re-simulate."'
    ))

    # ══ Glossary ══════════════════════════════════════════════════════════════
    _heading1(doc, "Glossary")
    _table(doc,
        ["Term", "Plain English Definition"],
        [
            ["Baseline",              "What you would have sold anyway, without the promotion running"],
            ["Incremental Units",     "The extra sales generated by the promotion above the baseline"],
            ["Lift",                  "The % increase in sales driven by the promo. 60% lift = selling 60% more than baseline"],
            ["Mechanic",              "The type of promotion: Price Off, BOGO (Buy-One-Get-One), Display, Feature, or Combo"],
            ["ROI",                   "Return on Investment — gross profit earned per rupee of trade spend invested"],
            ["Gross Profit",          "Revenue minus cost of goods. Does not include overheads or fixed costs"],
            ["Brand Investment",      "The cash you actually pay after the retailer contributes their co-funding share"],
            ["Co-funding",            "The portion of the discount funded by the retailer, not by your brand"],
            ["Cannibalization",       "When promoting Product A causes shoppers to switch away from Product B, reducing its sales"],
            ["Tier 1 Retailer",       "Your most important retail partners — highest volume, strategic accounts"],
            ["Compliance",            "Whether your plan follows all 17 internal trade policy rules"],
            ["Scenario",              "A saved version of a simulated plan. You can store and compare multiple scenarios"],
        ],
        col_widths=[4.5, 10.0],
    )

    doc.save("RGM_Business_Primer.docx")
    print("Saved: RGM_Business_Primer.docx")


# ── Entry point ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_technical_guide()
    build_business_primer()
    print("\nBoth documents generated successfully.")
